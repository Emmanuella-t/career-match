"""ATS-friendly resume export (DOCX and plain text)."""

from __future__ import annotations

import io
import re

from docx import Document

from career_match.tailoring.resume_structure import StructuredResume

_SECTION_TITLES = {
    "summary": "Summary",
    "experience": "Experience",
    "projects": "Projects",
    "skills": "Skills",
    "education": "Education",
    "other": "Additional",
}


def sanitize_export_filename(name: str, extension: str) -> str:
    """Return a safe download filename without path components."""
    base = re.sub(r"[^\w\s\-]", "", name).strip() or "resume"
    base = re.sub(r"\s+", "_", base)[:80]
    ext = extension.lstrip(".")
    return f"{base}_tailored.{ext}"


def export_plain_text(structured: StructuredResume) -> str:
    """Render structured resume as ATS-friendly plain text."""
    lines: list[str] = []
    if structured.header:
        lines.extend(structured.header.split("\n"))
        lines.append("")
    for block in structured.sections:
        if not block.text.strip():
            continue
        title = _SECTION_TITLES.get(block.section, block.section.title())
        lines.append(title.upper())
        lines.append(block.text.strip())
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def export_docx_bytes(structured: StructuredResume) -> bytes:
    """Build a single-column DOCX resume in memory."""
    document = Document()
    if structured.header:
        for line in structured.header.split("\n"):
            if line.strip():
                paragraph = document.add_paragraph(line.strip())
                if paragraph.runs:
                    paragraph.runs[0].bold = True

    for block in structured.sections:
        if not block.text.strip():
            continue
        title = _SECTION_TITLES.get(block.section, block.section.title())
        document.add_heading(title, level=1)
        for line in block.text.split("\n"):
            if line.strip():
                document.add_paragraph(line.strip())

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
