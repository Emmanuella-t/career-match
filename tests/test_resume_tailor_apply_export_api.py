"""Tests for tailoring apply, alignment comparison, and export."""

from __future__ import annotations

import io
from unittest.mock import patch

import numpy as np
import pytest
from docx import Document
from fastapi.testclient import TestClient

from career_match.api.app import create_app
from career_match.api.auth import ClerkIdentity
from career_match.api.services import MatcherService
from career_match.matching.hybrid import HybridMatcher
from career_match.matching.semantic import SemanticMatcher
from career_match.persistence.store import InMemoryPersistenceStore
from career_match.tailoring.export import export_docx_bytes, export_plain_text
from career_match.tailoring.protocol import RewriteSuggestion
from career_match.tailoring.providers import FakeRewriteProvider
from career_match.tailoring.resume_structure import apply_suggestions, parse_resume_sections
from career_match.tailoring.safeguards import validate_revised_resume
from career_match.tailoring.suggestion_ids import compute_suggestion_id

USER_A = ClerkIdentity(user_id="user_a", email="a@example.com", display_name="Ada")
USER_B = ClerkIdentity(user_id="user_b", email="b@example.com", display_name="Bea")

PYTHON_RESUME = (
    "Jordan Lee\nBackend Engineer\n\n"
    "Experience\n"
    "Built GitHub Actions pipelines for automated testing and deployment. "
    "Python, FastAPI, Docker, SQL, and Git on Linux.\n\n"
    "Skills: Python, FastAPI, Docker, SQL, Git"
)

CI_CD_JOB = (
    "Platform Engineer\n\n"
    "Looking for CI/CD experience and Python services on Linux."
)

AWS_JOB = (
    "Senior Backend Engineer\n\n"
    "Required: Python, FastAPI, Docker, AWS, CI/CD, Kubernetes.\n"
    "Must have production deployment experience."
)

GOOD_SUGGESTION = RewriteSuggestion(
    section="experience",
    original_text="Built GitHub Actions pipelines for automated testing and deployment.",
    suggested_text="Built GitHub Actions pipelines for CI/CD and automated testing and deployment.",
    keywords_introduced=("ci/cd",),
    support_reason="equivalent evidence",
    support_level="high",
)


class _FixedEncoder:
    def encode(self, texts):
        rows = []
        for text in texts:
            vec = np.zeros(4, dtype=np.float32)
            lowered = text.lower()
            vec[0] = 1.0 if "python" in lowered else 0.2
            vec[1] = 1.0 if "docker" in lowered else 0.1
            vec[2] = 1.0 if "ci/cd" in lowered or "github actions" in lowered else 0.2
            vec[3] = min(1.0, len(text) / 400.0)
            rows.append(vec)
        matrix = np.vstack(rows)
        norms = np.linalg.norm(matrix, axis=1, keepdims=True)
        return matrix / np.clip(norms, 1e-12, None)


@pytest.fixture()
def store() -> InMemoryPersistenceStore:
    return InMemoryPersistenceStore()


@pytest.fixture()
def client(store: InMemoryPersistenceStore) -> TestClient:
    semantic = SemanticMatcher(encoder=_FixedEncoder())
    hybrid = HybridMatcher(semantic_matcher=semantic)
    service = MatcherService(semantic=semantic, hybrid=hybrid)

    application = create_app()
    application.state.persistence_store = store
    application.state.clerk_identity_override = USER_A
    application.state.matcher_service = service

    with TestClient(application) as test_client:
        yield test_client


def _auth_headers() -> dict[str, str]:
    return {"Authorization": "Bearer test-token"}


def _create_resume(client: TestClient, text: str = PYTHON_RESUME) -> str:
    response = client.post(
        "/api/v1/resumes",
        headers=_auth_headers(),
        json={"name": "Primary", "resume_text": text},
    )
    assert response.status_code == 201
    return response.json()["id"]


def _tailor_with_fake(client: TestClient, resume_id: str, job: str = CI_CD_JOB) -> dict:
    client.app.state.rewrite_provider_override = FakeRewriteProvider((GOOD_SUGGESTION,))
    response = client.post(
        "/api/v1/resumes/tailor",
        headers=_auth_headers(),
        json={"resume_id": resume_id, "job_description": job, "matcher": "semantic"},
    )
    assert response.status_code == 200
    return response.json()


def test_suggestion_ids_are_stable() -> None:
    sid = compute_suggestion_id(GOOD_SUGGESTION)
    assert sid == compute_suggestion_id(GOOD_SUGGESTION)
    assert len(sid) == 16


def test_apply_accepts_one_supported_suggestion(client: TestClient) -> None:
    resume_id = _create_resume(client)
    tailor_body = _tailor_with_fake(client, resume_id)
    suggestion_id = tailor_body["rewrite_suggestions"][0]["suggestion_id"]

    response = client.post(
        "/api/v1/resumes/tailor/apply",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": [suggestion_id],
            "matcher": "semantic",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["accepted_suggestions"]
    assert body["rejected_suggestions"] == []
    assert "ci/cd" in body["revised_resume_text"].lower()
    assert body["alignment_delta"] == round(
        body["revised_alignment_score"] - body["original_alignment_score"],
        2,
    )
    assert body["revised_alignment_score"] >= body["original_alignment_score"]


def test_apply_rejects_unknown_suggestion_id(client: TestClient) -> None:
    resume_id = _create_resume(client)
    _tailor_with_fake(client, resume_id)

    response = client.post(
        "/api/v1/resumes/tailor/apply",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": ["deadbeefdeadbeef"],
        },
    )
    assert response.status_code == 400
    assert "unknown or stale suggestion" in response.json()["detail"].lower()


def test_rejected_suggestions_listed_when_not_accepted(client: TestClient) -> None:
    resume_id = _create_resume(client)
    tailor_body = _tailor_with_fake(client, resume_id)
    suggestion_id = tailor_body["rewrite_suggestions"][0]["suggestion_id"]

    response = client.post(
        "/api/v1/resumes/tailor/apply",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": [],
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["accepted_suggestions"] == []
    assert len(body["rejected_suggestions"]) == 1
    assert body["rejected_suggestions"][0]["suggestion_id"] == suggestion_id
    assert body["alignment_delta"] == 0.0
    assert "github actions" in body["revised_resume_text"].lower()


def test_tailor_filters_fabricated_suggestion(client: TestClient) -> None:
    fabricated = RewriteSuggestion(
        section="experience",
        original_text="Built GitHub Actions pipelines for automated testing and deployment.",
        suggested_text="Built AWS and Kubernetes platforms with 10 years experience.",
        keywords_introduced=("aws", "kubernetes"),
        support_reason="fake",
        support_level="low",
    )
    client.app.state.rewrite_provider_override = FakeRewriteProvider((fabricated,))
    resume_id = _create_resume(client)
    response = client.post(
        "/api/v1/resumes/tailor",
        headers=_auth_headers(),
        json={"resume_id": resume_id, "job_description": AWS_JOB},
    )
    body = response.json()
    assert body["rewrite_suggestions"] == []


def test_original_resume_not_overwritten_after_apply_and_export(client: TestClient) -> None:
    resume_id = _create_resume(client)
    tailor_body = _tailor_with_fake(client, resume_id)
    suggestion_id = tailor_body["rewrite_suggestions"][0]["suggestion_id"]

    client.post(
        "/api/v1/resumes/tailor/apply",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": [suggestion_id],
        },
    )
    client.post(
        "/api/v1/resumes/export",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": [suggestion_id],
            "format": "txt",
        },
    )

    stored = client.get(f"/api/v1/resumes/{resume_id}", headers=_auth_headers())
    assert stored.status_code == 200
    assert stored.json()["resume_text"] == PYTHON_RESUME


def test_remaining_gaps_retained_after_apply(client: TestClient) -> None:
    resume_id = _create_resume(client)
    tailor_body = _tailor_with_fake(client, resume_id, AWS_JOB)
    suggestion_id = tailor_body["rewrite_suggestions"][0]["suggestion_id"]

    response = client.post(
        "/api/v1/resumes/tailor/apply",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": AWS_JOB,
            "accepted_suggestion_ids": [suggestion_id],
        },
    )
    body = response.json()
    remaining = {item.lower() for item in body["remaining_unsupported_keywords"]}
    assert "aws" in remaining or "kubernetes" in remaining


def test_apply_cross_user_resume_rejected(client: TestClient) -> None:
    resume_id = _create_resume(client)
    client.app.state.clerk_identity_override = USER_B
    response = client.post(
        "/api/v1/resumes/tailor/apply",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": [],
        },
    )
    assert response.status_code == 404


def test_export_txt_valid(client: TestClient) -> None:
    resume_id = _create_resume(client)
    tailor_body = _tailor_with_fake(client, resume_id)
    suggestion_id = tailor_body["rewrite_suggestions"][0]["suggestion_id"]

    response = client.post(
        "/api/v1/resumes/export",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": [suggestion_id],
            "format": "txt",
        },
    )
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/plain")
    assert "Jordan Lee" in response.text
    assert "EXPERIENCE" in response.text.upper()
    disposition = response.headers.get("content-disposition", "")
    assert "attachment" in disposition.lower()
    assert ".." not in disposition


def test_export_docx_openable(client: TestClient) -> None:
    resume_id = _create_resume(client)
    tailor_body = _tailor_with_fake(client, resume_id)
    suggestion_id = tailor_body["rewrite_suggestions"][0]["suggestion_id"]

    response = client.post(
        "/api/v1/resumes/export",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "accepted_suggestion_ids": [suggestion_id],
            "format": "docx",
        },
    )
    assert response.status_code == 200
    document = Document(io.BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Jordan Lee" in text
    assert "Experience" in text or "EXPERIENCE" in text.upper()


def test_keyword_stuffing_blocks_apply(client: TestClient) -> None:
    resume_id = _create_resume(client)
    tailor_body = _tailor_with_fake(client, resume_id)
    suggestion_id = tailor_body["rewrite_suggestions"][0]["suggestion_id"]

    with patch(
        "career_match.api.tailor_apply_service.validate_revised_resume",
        return_value=["Revised resume appears keyword-stuffed"],
    ):
        response = client.post(
            "/api/v1/resumes/tailor/apply",
            headers=_auth_headers(),
            json={
                "resume_id": resume_id,
                "job_description": CI_CD_JOB,
                "accepted_suggestion_ids": [suggestion_id],
            },
        )
    assert response.status_code == 400
    assert "keyword-stuffed" in response.json()["detail"].lower()


def test_export_invalid_format_rejected(client: TestClient) -> None:
    resume_id = _create_resume(client)
    response = client.post(
        "/api/v1/resumes/export",
        headers=_auth_headers(),
        json={
            "resume_id": resume_id,
            "job_description": CI_CD_JOB,
            "format": "pdf",
        },
    )
    assert response.status_code == 422


def test_export_requires_auth(store: InMemoryPersistenceStore) -> None:
    application = create_app()
    application.state.persistence_store = store
    with TestClient(application) as test_client:
        response = test_client.post(
            "/api/v1/resumes/export",
            json={
                "resume_text": PYTHON_RESUME,
                "job_description": CI_CD_JOB,
                "format": "txt",
            },
        )
    assert response.status_code == 401


def test_structured_apply_tracks_accepted_block() -> None:
    structured = parse_resume_sections(PYTHON_RESUME)
    suggestion_id = compute_suggestion_id(GOOD_SUGGESTION)
    revised = apply_suggestions(structured, ((suggestion_id, GOOD_SUGGESTION),))
    experience = next(block for block in revised.sections if block.section == "experience")
    assert experience.change_status == "accepted"
    assert "CI/CD" in experience.text or "ci/cd" in experience.text.lower()


def test_export_helpers_produce_content() -> None:
    structured = parse_resume_sections(PYTHON_RESUME)
    plain = export_plain_text(structured)
    assert "Jordan Lee" in plain
    docx_bytes = export_docx_bytes(structured)
    document = Document(io.BytesIO(docx_bytes))
    assert document.paragraphs


def test_safeguards_flag_forbidden_metrics() -> None:
    original = "Built Python APIs."
    revised = "Built Python APIs with 10 years experience."
    warnings = validate_revised_resume(
        original_text=original,
        revised_text=revised,
        job_text="Need Python.",
        unsupported_keywords=(),
    )
    assert warnings
